"""DOCX document parser."""

import time
from typing import Dict, List

from docx import Document as DocxDocument

from ..base import BaseParser
from ..types import MossDocument, ParseResult


class DocxParser(BaseParser):
    """Parser for DOCX files."""

    def parse(self, file_path: str) -> ParseResult:
        """Parse a DOCX file and extract text from paragraphs.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            ParseResult containing one document per paragraph (or chunked if needed).
        """
        start_time = time.time()

        documents = []
        doc = DocxDocument(file_path)
        
        for para_num, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if text.strip():  # Only add non-empty paragraphs
                doc_id = f"{file_path}_para_{para_num}"
                metadata = {
                    "source_file": file_path,
                    "paragraph_number": para_num + 1,  # 1-indexed for humans
                    "total_paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                }
                documents.append(
                    MossDocument(
                        id=doc_id,
                        text=text.strip(),
                        metadata=metadata,
                    )
                )

        parse_time_ms = (time.time() - start_time) * 1000
        return ParseResult(
            documents=documents,
            source_path=file_path,
            parse_time_ms=parse_time_ms,
        )

    def supported_extensions(self) -> List[str]:
        """Return a list of file extensions this parser supports.

        Returns:
            List of file extensions (without the dot).
        """
        return ["docx"]